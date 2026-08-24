import io
import os
import base64
import re
import asyncio
from typing import Optional
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, Flowable
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
from fastapi.responses import StreamingResponse
from app.core.logging import get_logger
import pypdf

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from a PDF file using pypdf."""
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

logger = get_logger()

_META_EMPTY = "—"


def _normalize_section_key(value) -> str:
    return str(value or "").strip().lower().replace(":", "").replace(" ", "_")


def _show_weights(content: dict, section_key: str) -> bool:
    """Dynamically check if weights should be shown for any section based on weight_view_<section> lock."""
    if not isinstance(content, dict):
        return True
    
    norm = _normalize_section_key(section_key)
    
    # Check for weight_view_<section> lock dynamically
    possible_keys = [
        f"weight_view_{norm}",
        f"weight_view_{norm}_view",
        f"weight_view_{norm.replace('_', '')}",
        f"weight_view_{norm.replace('_', '')}_view",
    ]
    
    for k in possible_keys:
        if content.get(k) == "locked":
            return False
    
    # Default to showing weights if no lock is found
    return True


def _is_weighted_list(value) -> bool:
    """Check if value is a weighted list (array of objects with point/weight fields)."""
    if not isinstance(value, list):
        return False
    if not value:
        return False
    # Check if first item has point/weight structure
    first_item = value[0]
    if isinstance(first_item, dict):
        return any(k in first_item for k in ("point", "text", "title"))
    return False


def _format_section_title(key: str, labels: dict = None) -> str:
    """Convert a JSON key into a readable heading, using custom labels if provided."""
    norm = _normalize_section_key(key)
    if labels and isinstance(labels, dict):
        for l_key, l_val in labels.items():
            if _normalize_section_key(l_key) == norm:
                return str(l_val)
    return key.replace("_", " ").title()


def _is_section_visible(content: dict, section_key: str, sections_metadata: Optional[dict] = None) -> bool:
    """Check if a section should be visible based on metadata locks and view keys."""
    meta = sections_metadata or {}
    locks = meta.get("locks") or {}
    norm = _normalize_section_key(section_key)
    if locks.get(section_key) == "locked" or locks.get(norm) == "locked":
        return False
    view_key = f"{norm}_view"
    view_value = content.get(view_key)
    if view_value == "locked":
        return False
    return True


def _is_metadata_key(key: str) -> bool:
    """Check if a key is metadata (view locks, weight views, etc.) and should be skipped."""
    norm = _normalize_section_key(key)
    return (
        norm.endswith("_view")
        or norm.startswith("weight_view_")
        or norm.lstrip("_") in {"section_labels", "dynamic_sections", "section_visibility"})


def _sort_sections(content_dict: dict) -> list[tuple[str, any]]:
    """Sort sections based on specified section order, keeping custom sections at the end."""
    section_order = content_dict.get("_section_order")
    
    # Filter out metadata keys first
    filtered_items = [
        (k, v) for k, v in content_dict.items()
        if not _is_metadata_key(k) and k not in ("_section_order", "_section_labels", "_dynamic_sections")
    ]
    
    if isinstance(section_order, list):
        def get_order_index(item) -> int:
            norm_k = _normalize_section_key(item[0])
            for i, order_k in enumerate(section_order):
                if _normalize_section_key(order_k) == norm_k:
                    return i
            return len(section_order)
        return sorted(filtered_items, key=get_order_index)
        
    return filtered_items


def _export_str(v) -> str:
    if v is None:
        return ""
    text = str(v).strip()
    text = re.sub(r'\[\[mod:[^\]]+\]\]', '', text)
    text = re.sub(r'\[\[/mod\]\]', '', text)
    return text


def _export_first(data: dict, keys: list) -> str:
    for k in keys:
        v = data.get(k)
        if v is not None and (not isinstance(v, str) or v.strip()):
            text = str(v).strip()
            text = re.sub(r'\[\[mod:[^\]]+\]\]', '', text)
            text = re.sub(r'\[\[/mod\]\]', '', text)
            return text
    return ""


def _export_expand_salary_value(raw: str, unit: str) -> str:
    raw = str(raw).strip().replace(",", "")
    unit_lower = str(unit or "").strip().lower()
    try:
        val = float(raw)
        if unit_lower == "k":
            val *= 1_000
        elif unit_lower == "m":
            val *= 1_000_000
        return f"{int(val):,}"
    except ValueError:
        return raw


def _build_location_display_row(data: dict) -> str:
    loc = _export_str(data.get("location"))
    city = _export_str(data.get("city"))
    cc = _export_str(data.get("country_code"))
    parts: list[str] = []
    if loc:
        parts.append(loc)
    blob = " ".join(parts).lower()
    if city and city.lower() not in blob:
        parts.append(city)
        blob = " ".join(parts).lower()
    if cc and cc.lower() not in blob:
        parts.append(cc)
    return ", ".join(parts)


def _build_salary_display_row(data: dict) -> str:
    sr = _export_str(data.get("salary_range"))
    if sr:
        return sr
    symbol = _export_first(data, ["salary_symbol", "currency_symbol", "salary_currency_symbol"])
    min_val = _export_first(data, ["salary_min_value", "salary_min", "salary_value_min", "salary_from"])
    max_val = _export_first(data, ["salary_max_value", "salary_max", "salary_value_max", "salary_to"])
    unit = _export_first(data, ["salary_period", "salary_unit", "salary_unit_suffix"]) or ""
    if symbol and min_val and max_val:
        s = str(symbol).strip()
        mn = _export_expand_salary_value(str(min_val), unit)
        mx = _export_expand_salary_value(str(max_val), unit)
        return f"{s}{mn}-{s}{mx}"
    if symbol and min_val:
        s = str(symbol).strip()
        mn = _export_expand_salary_value(str(min_val), unit)
        return f"{s}{mn}+"
    return ""


def enrich_jd_export_dict(data: dict) -> dict:
    out = dict(data)
    loc_disp = _build_location_display_row(out)
    out["location_display"] = loc_disp or _export_str(out.get("location"))
    out["salary_display"] = _build_salary_display_row(out)
    return out


def _meta_value_for_pdf(safe_text_fn, raw) -> str:
    if raw is None or (isinstance(raw, str) and not str(raw).strip()):
        return _META_EMPTY
    text = str(raw)
    text = re.sub(r'\[\[mod:[^\]]+\]\]', '', text)
    text = re.sub(r'\[\[/mod\]\]', '', text)
    return safe_text_fn(text)


class TickMark(Flowable):
    def __init__(self, size=10, color="#10b981"):
        Flowable.__init__(self)
        self.size = size
        self.color = colors.HexColor(color) if isinstance(color, str) else color
        self.width = size
        self.height = size

    def draw(self):
        self.canv.saveState()
        self.canv.setStrokeColor(self.color)
        self.canv.setLineWidth(self.size / 6)
        self.canv.setLineCap(1)
        self.canv.line(0, self.size * 0.4, self.size * 0.35, 0)
        self.canv.line(self.size * 0.35, 0, self.size, self.size * 0.9)
        self.canv.restoreState()


class CheckedBox(Flowable):
    def __init__(self, size=12, color="#10b981"):
        Flowable.__init__(self)
        self.size = size
        self.color = colors.HexColor(color) if isinstance(color, str) else color
        self.width = size
        self.height = size

    def draw(self):
        s = self.size
        c = self.canv
        c.saveState()
        
        # Draw the outer box
        c.setStrokeColor(colors.black)
        c.setLineWidth(1.1)
        c.rect(0, 0, s, s, stroke=1, fill=0)
        
        # Draw the tick inside
        c.setStrokeColor(self.color)
        c.setLineWidth(s / 6)
        c.setLineCap(1)
        # Positioning the tick to look balanced inside the square
        c.line(s * 0.22, s * 0.45, s * 0.45, s * 0.22)
        c.line(s * 0.45, s * 0.22, s * 0.82, s * 0.8)
        
        c.restoreState()


class PersonIcon(Flowable):
    def __init__(self, size=44, bg_color="#eef2ff", icon_color="#4f46e5"):
        Flowable.__init__(self)
        self.size = size
        self.bg = colors.HexColor(bg_color) if isinstance(bg_color, str) else bg_color
        self.icon = colors.HexColor(icon_color) if isinstance(icon_color, str) else icon_color
        self.width = size
        self.height = size

    def draw(self):
        s = self.size
        c = self.canv
        c.saveState()
        c.setFillColor(self.bg)
        c.setStrokeColor(self.bg)
        c.roundRect(0, 0, s, s, s * 0.22, fill=1, stroke=0)
        c.setFillColor(self.icon)
        head_r = s * 0.20
        c.circle(s * 0.5, s * 0.64, head_r, fill=1, stroke=0)
        body_w = s * 0.50
        body_h = s * 0.24
        body_x = (s - body_w) / 2
        body_y = s * 0.10
        c.roundRect(body_x, body_y, body_w, body_h, body_h * 0.5, fill=1, stroke=0)
        c.restoreState()


class PDFGenerator:
    def __init__(self):
        pass

    @staticmethod
    def _safe_text(value) -> str:
        if not value:
            return ""
        text = str(value)
        text = re.sub(r'\[\[mod:[^\]]+\]\]', '', text)
        text = re.sub(r'\[\[/mod\]\]', '', text)
        replacements = {
            "\u2010": "-", "\u2011": "-", "\u2012": "-", "\u2013": "-",
            "\u2014": "-", "\u2015": "-", "\u2212": "-", "\u00ad": "",
            "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
            "\u00a0": " ", "\u2022": "•",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text

    @staticmethod
    def _as_weighted_points(value):
        if isinstance(value, list):
            points = []
            for item in value:
                if isinstance(item, dict):
                    point = str(item.get("point") or item.get("text") or item.get("title") or "").strip()
                    point = PDFGenerator._safe_text(point)
                    try:
                        weight = int(item.get("weight"))
                    except (TypeError, ValueError):
                        weight = None
                    if point:
                        points.append({"point": point, "weight": weight})
                else:
                    point = str(item).strip()
                    point = PDFGenerator._safe_text(point)
                    if point:
                        points.append({"point": point, "weight": None})
            return points
        if isinstance(value, str) and value.strip():
            return [{"point": PDFGenerator._safe_text(value.strip()), "weight": None}]
        return []

    def _render_pdf_bytes(self, jd_data: dict, jd_title: str, exclude_terms: bool = False) -> bytes:
        jd_data = enrich_jd_export_dict(dict(jd_data))
        buffer = io.BytesIO()

        doc = SimpleDocTemplate(buffer,pagesize=letter,leftMargin=40,rightMargin=40,topMargin=35,bottomMargin=35)
        styles = getSampleStyleSheet()
        page_width = letter[0] - 80  # 532pt

        elements = []
        logger.debug("PDF render: starting for title=%s, content keys=%s", jd_title, list(jd_data.get("content", {}).keys()) if jd_data.get("content") else [])

        primary_blue  = colors.HexColor("#3b82f6")
        text_dark     = colors.HexColor("#1e293b")
        text_gray     = colors.HexColor("#64748b")
        border_color  = colors.HexColor("#e2e8f0")
        card_bg       = colors.HexColor("#f8fafc")

        try:
            pdfmetrics.registerFont(TTFont('Arial',      'C:/Windows/Fonts/arial.ttf'))
            pdfmetrics.registerFont(TTFont('Arial-Bold', 'C:/Windows/Fonts/arialbd.ttf'))
            font_path = os.path.join(os.getcwd(), "static", "fonts", "DancingScript-Regular.ttf")
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('DancingScript', font_path))
                signature_font = "DancingScript"
            else:
                signature_font = "Helvetica-Bold"
            main_font = "Arial"
            bold_font = "Arial-Bold"
        except Exception:
            main_font      = "Helvetica"
            bold_font      = "Helvetica-Bold"
            signature_font = "Helvetica-Bold"

        job_title_style = ParagraphStyle(
            "JobTitle", parent=styles["Heading1"],
            fontSize=20, fontName=bold_font, textColor=text_dark,
            spaceAfter=12, leading=24,
        )
        section_label_style = ParagraphStyle(
            "MetaLabel", parent=styles["Normal"],
            fontSize=7, fontName=bold_font, textColor=text_gray, leading=8,
        )
        section_value_style = ParagraphStyle(
            "MetaValue", parent=styles["Normal"],
            fontSize=9, fontName=bold_font, textColor=text_dark, leading=11,
        )
        body_style = ParagraphStyle(
            "Body", parent=styles["Normal"],
            fontSize=9, fontName=main_font, leading=14,
            textColor=colors.HexColor("#334155"),
        )
        section_header_style = ParagraphStyle(
            "SectionHeader", parent=styles["Heading2"],
            fontSize=11, fontName=bold_font, textColor=text_dark, leading=14,
        )
        signature_text_style = ParagraphStyle(
            "SignatureText", parent=styles["Normal"],
            fontSize=22, fontName=signature_font,
            textColor=colors.HexColor("#312e81"),
            leading=28, alignment=1,
        )
        verified_style = ParagraphStyle(
            "VerifiedText", parent=styles["Normal"],
            fontSize=8, fontName=bold_font,
            textColor=colors.HexColor("#10b981"),
            leading=10, alignment=1, spaceBefore=2,
        )
        signee_name_style = ParagraphStyle(
            "SigneeName", parent=styles["Normal"],
            fontSize=11, fontName=bold_font,
            textColor=colors.HexColor("#1e293b"),
            leading=14, spaceBefore=2,
        )
        signee_role_style = ParagraphStyle(
            "SigneeRole", parent=styles["Normal"],
            fontSize=9, fontName=main_font,
            textColor=colors.HexColor("#64748b"),
            leading=11, spaceBefore=1,
        )
        bold_body_style = ParagraphStyle(
            "BoldBody", parent=body_style,
            fontName=bold_font,
        )

        # Handle image(s) for PDF – support a single URL string or a list of URLs
        image_urls = jd_data.get("image_url")
        # Normalize to a list
        if isinstance(image_urls, str):
            image_urls = [image_urls]
        if isinstance(image_urls, list):
            for image_url in image_urls:
                if not image_url:
                    continue
                try:
                    if isinstance(image_url, str) and (image_url.startswith("http://") or image_url.startswith("https://")):
                        import requests
                        response = requests.get(image_url, timeout=5)
                        response.raise_for_status()
                        logo = RLImage(io.BytesIO(response.content))
                    else:
                        rel_path = str(image_url or "").lstrip("/")
                        abs_path = os.path.join(os.getcwd(), rel_path)
                        if os.path.isfile(abs_path):
                            logo = RLImage(abs_path)
                        else:
                            continue
                    fixed_w = 380
                    ratio = fixed_w / logo.drawWidth
                    logo.drawWidth = fixed_w
                    logo.drawHeight *= ratio
                    logo.hAlign = "CENTER"
                    elements.append(logo)
                    elements.append(Spacer(1, 10))
                except Exception as e:
                    logger.debug(f"Failed to load image {image_url}: {str(e)}")
                    pass

        elements.append(Paragraph(self._safe_text(jd_title), job_title_style))

        meta_fields = [
            ("JOB ID",           "job_id"),
            ("DEPARTMENT",       "department"),
            ("JOB FAMILY",       "job_family"),
            ("INDUSTRY",         "industry"),
            ("LOCATION",         "location_display"),
            ("JOB LEVEL",        "job_level"),
            ("SENIORITY",        "seniority"),
            ("EMPLOYMENT TYPE",  "employment_type"),
            ("SALARY RANGE",     "salary_display"),
        ]
        filled = [
            (lbl, _meta_value_for_pdf(self._safe_text, jd_data.get(key)))
            for lbl, key in meta_fields
        ]
        if filled:
            while len(filled) % 4 != 0:
                filled.append(("", ""))
            rows = []
            col_w = page_width / 4
            for i in range(0, len(filled), 4):
                label_row = [Paragraph(lbl, section_label_style) for lbl, _ in filled[i:i+4]]
                value_row = [Paragraph(val, section_value_style) for _, val in filled[i:i+4]]
                rows.append(label_row)
                rows.append(value_row)
                rows.append([Spacer(1, 4)] * 4)
            meta_tbl = Table(rows, colWidths=[col_w] * 4)
            meta_tbl.setStyle(TableStyle([
                ("VALIGN",       (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING",   (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
                ("LEFTPADDING",  (0, 0), (-1, -1), 8),
                ("BACKGROUND",   (0, 0), (-1, -1), card_bg),
                ("BOX",          (0, 0), (-1, -1), 0.5, border_color),
                ("LINEBELOW",    (0, 0), (-1, -1), 0.5, border_color),
            ]))
            elements.append(meta_tbl)

        elements.append(Spacer(1, 15))

        def add_card_header(title: str, accent_color=primary_blue):
            elements.append(Paragraph(title.upper(), section_header_style))
            tbl = Table([[""]], colWidths=[page_width], rowHeights=[1.5],
                        style=TableStyle([
                            ("BACKGROUND", (0, 0), (-1, -1), accent_color),
                            ("TOPPADDING", (0, 0), (-1, -1), 0),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                            ("LEFTPADDING", (0, 0), (-1, -1), 0),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ]))
            elements.append(tbl)
            elements.append(Spacer(1, 8))

        def add_weighted_content(items, show_weights: bool = True):
            if not items: return
            rows = [[
                Paragraph("<b>DESCRIPTION</b>", section_label_style),
                Paragraph("<b>WEIGHT (%)</b>",  section_label_style),
            ]]
            if not show_weights:
                rows = [[Paragraph("<b>DESCRIPTION</b>", section_label_style)]]
            for i in items:
                w = f"{i.get('weight')}%" if i.get("weight") is not None else "—"
                row = [Paragraph(self._safe_text(i.get("point")), body_style)]
                if show_weights:
                    row.append(Paragraph(w, body_style))
                rows.append(row)
            t = Table(rows, colWidths=([page_width - 80, 80] if show_weights else [page_width]))
            t.setStyle(TableStyle([
                ("LINEBELOW",    (0, 0), (-1, -1), 0.2, border_color),
                ("TOPPADDING",   (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 6),
                ("VALIGN",       (0, 0), (-1, -1), "TOP"),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 12))

        content = jd_data.get("content", {})
        sections_metadata = jd_data.get("sections_metadata") or {}
        logger.debug(f"PDF render: All content keys before filtering: {list(content.keys())}")

        # Completely dynamic section rendering - no hardcoded section names
        # Color palette for dynamic sections
        section_colors = [
            "#3b82f6",  # blue
            "#6366f1",  # indigo
            "#10b981",  # green
            "#f59e0b",  # amber
            "#8b5cf6",  # purple
            "#ef4444",  # red
            "#f97316",  # orange
            "#0ea5e9",  # sky
            "#ec4899",  # pink
            "#14b8a6",  # teal
            "#4f46e5",  # royal indigo
            "#06b6d4",  # cyan
            "#059669",  # emerald
            "#d97706",  # dark amber
            "#7c3aed",  # violet
            "#db2777",  # dark pink
            "#2563eb",  # dark blue
            "#0284c7",  # light sky
            "#0d9488",  # dark teal
            "#4338ca",  # deep indigo
        ]
        color_index = 0

        for key, value in _sort_sections(content):
            # Skip metadata keys (view locks, weight views)
            if _is_metadata_key(key):
                logger.debug(f"PDF render: Skipping metadata key: {key}")
                continue
            
            # Skip empty values
            if not value:
                logger.debug(f"PDF render: Skipping empty value for key: {key}")
                continue
            
            # Check visibility via sections_metadata locks and legacy view keys
            if not _is_section_visible(content, key, sections_metadata):
                logger.debug(f"PDF render: Section {key} is not visible (locked)")
                continue

            logger.debug(f"PDF render: Rendering section: {key}, type: {type(value)}")

            # Format the title dynamically from the key
            title = _format_section_title(key, content.get("_section_labels"))
            
            # Get color cyclically
            color_hex = section_colors[color_index % len(section_colors)]
            color_index += 1

            # Render based on data type, not key name
            if _is_weighted_list(value):
                # Weighted list section (array of {point, weight} objects)
                add_card_header(title, colors.HexColor(color_hex))
                add_weighted_content(self._as_weighted_points(value), _show_weights(content, key))
            elif isinstance(value, list):
                # Plain list (not weighted) - render as bullet points
                add_card_header(title, colors.HexColor(color_hex))
                # Convert plain list to weighted points without weights
                plain_points = [{"point": str(item).strip(), "weight": None} for item in value if str(item).strip()]
                add_weighted_content(plain_points, False)  # Never show weights for plain lists
            else:
                # Text section (string or other scalar)
                text_value = self._safe_text(value)
                
                # Apply salary formatting to any text section that might contain salary references
                # This is a generic heuristic, not tied to specific section names
                salary = jd_data.get("salary_range", "")
                if salary and "k" in salary.lower():
                    matches = re.findall(r"\$(\d+)k?\s*-\s*\$?(\d+)k?", salary.lower())
                    if matches:
                        s_min, s_max = matches[0]
                        text_value = text_value.replace(f"${s_min}-${s_max}", f"${s_min}k-${s_max}k")
                        text_value = text_value.replace(f"${s_min} - ${s_max}", f"${s_min}k - ${s_max}k")
                
                add_card_header(title, colors.HexColor(color_hex))
                elements.append(Paragraph(text_value, body_style))
                elements.append(Spacer(1, 12))

        if not exclude_terms:
            # Add the Terms and Conditions section header
            elements.append(Spacer(1, 12))
            add_card_header("Terms and Conditions", colors.HexColor("#475569"))
            
            # Add the static organizational adherence clause with a ticked checkbox
            elements.append(Spacer(1, 12))
            terms_text = "By accepting this position, the employee agrees to adhere to all organizational policies, confidentiality obligations, professional standards, and applicable employment terms associated with the role."
            
            # Format terms for PDF: Bold headings
            formatted_terms = []
            for line in terms_text.split('\n'):
                trimmed = line.strip()
                if not trimmed: continue
                # Detect headings (similar logic to frontend)
                is_heading = any(trimmed.lower().endswith(w) for w in ["clause", "statement", "policy", "notice", "agreement", "acknowledgement"])
                if is_heading and len(trimmed.split()) <= 6:
                    formatted_terms.append(f"<b>{trimmed.upper()}</b>")
                else:
                    formatted_terms.append(trimmed)
            
            adherence_html = "<br/>".join(formatted_terms)
            
            # Use CheckedBox and bold body style for the terms
            elements.append(Table([
                [CheckedBox(11), Paragraph(adherence_html, body_style)]
            ], colWidths=[22, page_width - 22], style=TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ])))

        if jd_data.get("terms_accepted") and not exclude_terms:
            elements.append(Spacer(1, 20))
            elements.append(Table([
                [TickMark(12), Paragraph("<b>Accept Terms & Conditions</b>", section_header_style)],
                ["", Paragraph(
                    "I accept and agree that all information provided in this Job Description is accurate "
                    "and complete to the best of my knowledge. I acknowledge that the roles, responsibilities, "
                    "qualifications, and company details mentioned in this JD comply with organizational hiring "
                    "policies and applicable regulations. I also understand and consent that this Job Description "
                    "may be processed, stored, and shared internally for recruitment and evaluation purposes.",
                    body_style,
                )],
            ], colWidths=[20, page_width - 20], style=TableStyle([
                ("VALIGN",       (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING",   (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING",(0, 0), (-1,-1),  2),
            ])))

        # Render Signature block if JD was accepted/signed
        is_signed_off = jd_data.get("terms_accepted") or (jd_data.get("completed_at") is not None)
        if is_signed_off:
            elements.append(Spacer(1, 20))

            elements.append(Paragraph("DIGITAL SIGNATURE DETAILS", section_header_style))
            sig_line = Table([[""]], colWidths=[page_width], rowHeights=[1.5],
                             style=TableStyle([
                                 ("BACKGROUND", (0,0), (-1,-1), primary_blue),
                                 ("TOPPADDING", (0,0), (-1,-1), 0),
                                 ("BOTTOMPADDING", (0,0), (-1,-1), 0),
                                 ("LEFTPADDING", (0,0), (-1,-1), 0),
                                 ("RIGHTPADDING", (0,0), (-1,-1), 0),
                             ]))
            elements.append(sig_line)
            elements.append(Spacer(1, 14))

            gap    = 20
            card_w = (page_width - gap) / 2
            CARD_H = 120
            PAD    = 18

            sig_method = jd_data.get("signature_method")
            # If accepted with password, use username logic as requested by user
            if sig_method == "password":
                signee_name = self._safe_text(jd_data.get("username") or jd_data.get("signed_by") or "User")
            else:
                signee_name = self._safe_text(jd_data.get("candidate_name") or jd_data.get("signed_by") or "Candidate")
            
            signee_role = self._safe_text(jd_data.get("candidate_role") or "Employee")

            checkbox_style = ParagraphStyle(
                "CheckboxBadge", parent=styles["Normal"],
                fontSize=7, fontName=bold_font,
                textColor=colors.HexColor("#10b981"),
                leading=9,
            )

            badge_tbl = Table(
                [[TickMark(8), Paragraph("SIGNEE", checkbox_style)]],
                colWidths=[12, 55],
                style=TableStyle([
                    ("VALIGN",      (0,0), (-1,-1), "MIDDLE"),
                    ("LEFTPADDING", (0,0), (0,0),   0),
                    ("LEFTPADDING", (0,0), (1,0),   3),
                    ("TOPPADDING",  (0,0), (-1,-1),  0),
                    ("BOTTOMPADDING",(0,0),(-1,-1),  0),
                ]),
            )

            icon_size = 44
            icon_tbl = Table(
                [[PersonIcon(icon_size)]],
                colWidths=[icon_size], rowHeights=[icon_size],
            )

            info_inner_w = card_w - (PAD * 2) - icon_size - 8
            info_tbl = Table(
                [
                    [Paragraph(signee_name, signee_name_style)],
                    [Paragraph(signee_role, signee_role_style)],
                ],
                colWidths=[info_inner_w],
            )
            info_tbl.setStyle(TableStyle([
                ("TOPPADDING",   (0,0), (-1,-1), 0),
                ("BOTTOMPADDING",(0,0), (-1,-1), 0),
                ("LEFTPADDING",  (0,0), (-1,-1), 0),
            ]))

            icon_info_tbl = Table(
                [[icon_tbl, info_tbl]],
                colWidths=[icon_size + 8, info_inner_w],
            )
            icon_info_tbl.setStyle(TableStyle([
                ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
                ("LEFTPADDING",  (0,0), (0,0),   0),
                ("LEFTPADDING",  (0,0), (1,0),   8),
                ("TOPPADDING",   (0,0), (-1,-1),  0),
                ("BOTTOMPADDING",(0,0), (-1,-1),  0),
            ]))

            left_inner = Table(
                [
                    [badge_tbl],
                    [Spacer(1, 8)],
                    [icon_info_tbl],
                ],
                colWidths=[card_w - PAD * 2],
            )
            left_inner.setStyle(TableStyle([
                ("TOPPADDING",   (0,0), (-1,-1), 0),
                ("BOTTOMPADDING",(0,0), (-1,-1), 0),
                ("LEFTPADDING",  (0,0), (-1,-1), 0),
            ]))

            left_card = Table(
                [[left_inner]],
                colWidths=[card_w],
                rowHeights=[CARD_H],
            )
            left_card.setStyle(TableStyle([
                ("BACKGROUND",   (0,0), (-1,-1), card_bg),
                ("BOX",          (0,0), (-1,-1), 0.5, border_color),
                ("LEFTPADDING",  (0,0), (-1,-1), PAD),
                ("RIGHTPADDING", (0,0), (-1,-1), PAD),
                ("TOPPADDING",   (0,0), (-1,-1), PAD),
                ("BOTTOMPADDING",(0,0), (-1,-1), PAD),
                ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
            ]))

            sig_url = jd_data.get("signature_image_url") or jd_data.get("digital_signature_url")
            sig_content = None

            if sig_url:
                if str(sig_url).startswith("data:image"):
                    try:
                        # Handle base64 image data
                        if "," in str(sig_url):
                            header, encoded = str(sig_url).split(",", 1)
                            img_data = base64.b64decode(encoded)
                            sig_img = RLImage(io.BytesIO(img_data))
                            
                            max_sig_w = card_w - PAD * 2
                            max_sig_h = 56
                            aspect = sig_img.drawHeight / sig_img.drawWidth
                            if sig_img.drawWidth > max_sig_w:
                                sig_img.drawWidth = max_sig_w
                                sig_img.drawHeight = max_sig_w * aspect
                            if sig_img.drawHeight > max_sig_h:
                                sig_img.drawHeight = max_sig_h
                                sig_img.drawWidth = max_sig_h / aspect
                            sig_img.hAlign = "CENTER"
                            sig_content = sig_img
                    except Exception as e:
                        logger.error(f"Error decoding base64 signature: {e}")
                else:
                    # Handle file path
                    rel_path = sig_url.lstrip("/")
                    abs_sig_path = os.path.join(os.getcwd(), rel_path)
                    if os.path.isfile(abs_sig_path):
                        try:
                            sig_img = RLImage(abs_sig_path)
                            max_sig_w = card_w - PAD * 2
                            max_sig_h = 56
                            aspect = sig_img.drawHeight / sig_img.drawWidth
                            if sig_img.drawWidth > max_sig_w:
                                sig_img.drawWidth = max_sig_w
                                sig_img.drawHeight = max_sig_w * aspect
                            if sig_img.drawHeight > max_sig_h:
                                sig_img.drawHeight = max_sig_h
                                sig_img.drawWidth = max_sig_h / aspect
                            sig_img.hAlign = "CENTER"
                            sig_content = sig_img
                        except Exception as e:
                            logger.error(f"Error loading signature file: {e}")

            if sig_content is None:
                sig_content = Paragraph(signee_name, signature_text_style)

            verified_style_inline = ParagraphStyle(
                "VerifiedInline", parent=verified_style,
                alignment=0,  # left-align inside its tight column
                spaceBefore=0,
            )
            verified_badge = Table(
                [[TickMark(9), Paragraph("ELECTRONICALLY VERIFIED", verified_style_inline)]],
                colWidths=[13, 132],   # fixed narrow widths so tick hugs the text
                hAlign="CENTER",       # center the whole mini-table within the card
                style=TableStyle([
                    ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
                    ("LEFTPADDING",   (0,0), (0,0),   0),
                    ("RIGHTPADDING",  (0,0), (0,0),   0),
                    ("LEFTPADDING",   (0,0), (1,0),   4),
                    ("RIGHTPADDING",  (0,0), (1,0),   0),
                    ("TOPPADDING",    (0,0), (-1,-1),  0),
                    ("BOTTOMPADDING", (0,0), (-1,-1),  0),
                ]),
            )

            right_inner = Table(
                [
                    [sig_content],
                    [Spacer(1, 6)],
                    [verified_badge],
                ],
                colWidths=[card_w - PAD * 2],
            )
            right_inner.setStyle(TableStyle([
                ("ALIGN",        (0,0), (-1,-1), "CENTER"),
                ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
                ("TOPPADDING",   (0,0), (-1,-1),  0),
                ("BOTTOMPADDING",(0,0), (-1,-1),  0),
                ("LEFTPADDING",  (0,0), (-1,-1),  0),
            ]))

            right_card = Table(
                [[right_inner]],
                colWidths=[card_w],
                rowHeights=[CARD_H],
            )
            right_card.setStyle(TableStyle([
                ("BACKGROUND",   (0,0), (-1,-1), card_bg),
                ("BOX",          (0,0), (-1,-1), 0.5, border_color),
                ("LEFTPADDING",  (0,0), (-1,-1), PAD),
                ("RIGHTPADDING", (0,0), (-1,-1), PAD),
                ("TOPPADDING",   (0,0), (-1,-1), PAD),
                ("BOTTOMPADDING",(0,0), (-1,-1), PAD),
                ("ALIGN",        (0,0), (-1,-1), "CENTER"),
                ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
            ]))

            sig_outer_tbl = Table(
                [[left_card, Spacer(gap, 1), right_card]],
                colWidths=[card_w, gap, card_w],
                rowHeights=[CARD_H],
            )
            sig_outer_tbl.setStyle(TableStyle([
                ("VALIGN",       (0,0), (-1,-1), "TOP"),
                ("LEFTPADDING",  (0,0), (-1,-1), 0),
                ("RIGHTPADDING", (0,0), (-1,-1), 0),
                ("TOPPADDING",   (0,0), (-1,-1), 0),
                ("BOTTOMPADDING",(0,0), (-1,-1), 0),
            ]))
            elements.append(sig_outer_tbl)

        logger.debug("PDF render: building doc with %d elements", len(elements))
        doc.build(elements)
        buffer.seek(0)
        result = buffer.getvalue()
        logger.debug("PDF render: completed, returning %d bytes", len(result))
        return result

    async def generate_pdf_stream(self, jd_data: dict, jd_title: str, exclude_terms: bool = False):
        try:
            if not jd_data:
                logger.error("PDF export failed: jd_data is None or empty")
                raise ValueError("JD data is empty")
            
            if not jd_title:
                logger.warning("PDF export: jd_title is empty, using default")
            
            content = jd_data.get("content", {})
            if not content:
                logger.warning("PDF export: content dict is empty for title=%s", jd_title)
            else:
                logger.debug("PDF export: content keys = %s", list(content.keys()))
            
            pdf_bytes = await asyncio.to_thread(self._render_pdf_bytes, jd_data, jd_title, exclude_terms)
            if not pdf_bytes:
                logger.error("PDF export failed: _render_pdf_bytes returned empty bytes")
                raise ValueError("PDF generation produced no bytes")
            
            logger.debug("PDF export: generated %d bytes for title=%s", len(pdf_bytes), jd_title)
            
            role  = (jd_data.get("job_level") or jd_data.get("seniority") or "role")
            role  = role.replace(" ", "_") if role else "role"
            
            signed_by = jd_data.get("signed_by")
            filename = self._make_attachment_filename(jd_title or "job_description", role, signed_by, "pdf")
            logger.debug("PDF export: filename=%s", filename)
            
            return StreamingResponse(io.BytesIO(pdf_bytes),media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=\"{filename}\""})
        except Exception as e:
            logger.error(f"PDF generation failed: {type(e).__name__}: {str(e)}", exc_info=True)
            raise

    @staticmethod
    def _sanitize_filename(value: str) -> str:
        sanitized = re.sub(r'[^A-Za-z0-9_.-]+', '_', value or '')
        sanitized = re.sub(r'_{2,}', '_', sanitized).strip('_')
        return sanitized or 'job_description'

    def _make_attachment_filename(self, title: str, role: str, signed_by: str | None, ext: str) -> str:
        safe_title = self._sanitize_filename(title)
        safe_role = self._sanitize_filename(role)
        if signed_by:
            safe_signed_by = self._sanitize_filename(signed_by)
            filename = f"{safe_title}_{safe_role}_signed_by_{safe_signed_by}.{ext}"
        else:
            filename = f"{safe_title}_{safe_role}.{ext}"
        return filename

    def _render_word_bytes(self, jd_data: dict, jd_title: str, exclude_terms: bool = False) -> bytes:
        jd_data = enrich_jd_export_dict(dict(jd_data))
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor, Inches

        def set_bg(cell, hex_color: str):
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), hex_color)
            tc_pr.append(shd)

        doc = Document()

        section = doc.sections[0]
        section.top_margin    = Inches(0.7)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

        # Handle image(s) for Word – support a single URL string or a list of URLs
        image_urls = jd_data.get("image_url")
        if isinstance(image_urls, str):
            image_urls = [image_urls]
        if isinstance(image_urls, list):
            for image_url in image_urls:
                if not image_url:
                    continue
                try:
                    if isinstance(image_url, str) and (image_url.startswith("http://") or image_url.startswith("https://")):
                        import requests
                        response = requests.get(image_url, timeout=5)
                        response.raise_for_status()
                        logo_para = doc.add_paragraph()
                        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        logo_para.add_run().add_picture(io.BytesIO(response.content), width=Inches(6.25))
                        doc.add_paragraph()
                    else:
                        rel_path = str(image_url or "").lstrip("/")
                        abs_path = os.path.join(os.getcwd(), rel_path)
                        if os.path.isfile(abs_path):
                            logo_para = doc.add_paragraph()
                            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            logo_para.add_run().add_picture(abs_path, width=Inches(6.25))
                            doc.add_paragraph()
                except Exception as e:
                    logger.debug(f"Failed to load image {image_url}: {str(e)}")
                    pass

        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = t.add_run(jd_title)
        run.font.size = Pt(22)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
        p_pr  = t._p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'cbd5e1')
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
        doc.add_paragraph()

        meta_fields = [
            ("JOB ID",           "job_id"),
            ("DEPARTMENT",       "department"),
            ("JOB FAMILY",       "job_family"),
            ("INDUSTRY",         "industry"),
            ("LOCATION",         "location_display"),
            ("JOB LEVEL",        "job_level"),
            ("SENIORITY",        "seniority"),
            ("EMPLOYMENT TYPE",  "employment_type"),
            ("SALARY RANGE",     "salary_display"),
        ]
        filled = [
            (lbl, _meta_value_for_pdf(lambda t: t, jd_data.get(key)))
            for lbl, key in meta_fields
        ]
        if filled:
            while len(filled) % 4 != 0:
                filled.append(("", ""))
            rows_count = (len(filled) // 4) * 2
            meta_tbl = doc.add_table(rows=rows_count, cols=4)
            meta_tbl.style = None
            for idx, (lbl, val) in enumerate(filled):
                r, c = (idx // 4) * 2, idx % 4
                cell_lbl = meta_tbl.rows[r].cells[c]
                p_lbl = cell_lbl.paragraphs[0]
                p_lbl.space_before = Pt(4)
                run_l = p_lbl.add_run(lbl)
                run_l.font.size, run_l.font.color.rgb, run_l.bold = Pt(7), RGBColor(0x64, 0x74, 0x8b), True
                cell_val = meta_tbl.rows[r+1].cells[c]
                p_val = cell_val.paragraphs[0]
                p_val.space_after = Pt(6)
                run_v = p_val.add_run(val)
                run_v.font.size, run_v.bold = Pt(9.5), True
            for row in meta_tbl.rows:
                for cell in row.cells:
                    set_bg(cell, "f8fafc")
        doc.add_paragraph()

        content = jd_data.get("content", {})
        sections_metadata = jd_data.get("sections_metadata") or {}

        # Completely dynamic section rendering - no hardcoded section names
        # Color palette for dynamic sections
        section_colors = [
            "3b82f6",  # blue
            "6366f1",  # indigo
            "10b981",  # green
            "f59e0b",  # amber
            "8b5cf6",  # purple
            "ef4444",  # red
            "f97316",  # orange
            "0ea5e9",  # sky
            "ec4899",  # pink
            "14b8a6",  # teal
        ]
        color_index = 0

        def add_styled_section(title, text_val, color_hex, section_key=None):
            if not text_val: return
            show_weights = _show_weights(content, section_key or title)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
            run = p.add_run(title.upper())
            run.bold, run.font.size, run.font.color.rgb = True, Pt(11.5), RGBColor(0x1e, 0x29, 0x3b)
            p_pr  = p._p.get_or_add_pPr()
            p_bdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'),   'single')
            bottom.set(qn('w:sz'),    '12')  # 1.5pt thickness
            bottom.set(qn('w:space'), '2')   # 2pt space between text and line
            bottom.set(qn('w:color'), color_hex)
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
            doc.add_paragraph()

            if isinstance(text_val, list):
                items   = self._as_weighted_points(text_val)
                col_count = 2 if show_weights else 1
                w_table = doc.add_table(rows=0, cols=col_count)
                w_table.style   = None
                w_table.autofit = False
                w_table.columns[0].width = Inches(5.0)
                if show_weights:
                    w_table.columns[1].width = Inches(1.0)
                hdr_row = w_table.add_row().cells
                p_hdr_desc  = hdr_row[0].paragraphs[0]
                run_h_desc  = p_hdr_desc.add_run("DESCRIPTION")
                run_h_desc.font.size, run_h_desc.font.color.rgb, run_h_desc.bold = Pt(7), RGBColor(0x64, 0x74, 0x8b), True
                if show_weights:
                    p_hdr_w     = hdr_row[1].paragraphs[0]
                    p_hdr_w.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run_h_w     = p_hdr_w.add_run("WEIGHT (%)")
                    run_h_w.font.size, run_h_w.font.color.rgb, run_h_w.bold = Pt(7), RGBColor(0x64, 0x74, 0x8b), True
                for item in items:
                    row    = w_table.add_row().cells
                    p_text = row[0].paragraphs[0]
                    p_text.style = 'List Bullet'
                    point_text   = str(item.get("point") or "")
                    point_text   = re.sub(r'\[\[mod:[^\]]+\]\]', '', point_text)
                    point_text   = re.sub(r'\[\[/mod\]\]', '', point_text)
                    run_txt      = p_text.add_run(point_text)
                    run_txt.font.size = Pt(9.5)
                    if show_weights and item.get("weight") is not None:
                        p_weight = row[1].paragraphs[0]
                        p_weight.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        run_w = p_weight.add_run(f"({item.get('weight')}%)")
                        run_w.bold, run_w.font.size = True, Pt(9.5)
                        run_w.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            else:
                body_text = str(text_val)
                body_text = re.sub(r'\[\[mod:[^\]]+\]\]', '', body_text)
                body_text = re.sub(r'\[\[/mod\]\]', '', body_text)
                # Generic salary formatting - not tied to specific section names
                salary    = jd_data.get("salary_range", "")
                if salary and "k" in salary.lower():
                    matches = re.findall(r"\$(\d+)k?\s*-\s*\$?(\d+)k?", salary.lower())
                    if matches:
                        s_min, s_max = matches[0]
                        body_text = body_text.replace(f"${s_min}-${s_max}", f"${s_min}k-${s_max}k")
                        body_text = body_text.replace(f"${s_min} - ${s_max}", f"${s_min}k - ${s_max}k")
                p_body = doc.add_paragraph(body_text)
                for run in p_body.runs:
                    run.font.size = Pt(9.5)
            doc.add_paragraph()

        # Render all dynamic sections based on data type, not key names
        for key, value in _sort_sections(content):
            # Skip metadata keys (view locks, weight views)
            if _is_metadata_key(key):
                continue
            
            # Skip empty values
            if not value:
                continue
            
            # Check visibility via sections_metadata locks and legacy view keys
            if not _is_section_visible(content, key, sections_metadata):
                continue

            # Format the title dynamically from the key
            title = _format_section_title(key)
            
            # Get color cyclically
            color_hex = section_colors[color_index % len(section_colors)]
            color_index += 1

            # Render based on data type, not key name
            if _is_weighted_list(value):
                # Weighted list section
                add_styled_section(title, value, color_hex, key)
            elif isinstance(value, list):
                # Plain list - convert to weighted points without weights
                plain_points = [{"point": str(item).strip(), "weight": None} for item in value if str(item).strip()]
                add_styled_section(title, plain_points, color_hex, key)
            else:
                # Text section
                add_styled_section(title, value, color_hex, key)

        if jd_data.get("terms_accepted") and not exclude_terms:
            doc.add_paragraph()
            tc_p  = doc.add_paragraph()
            tc_p.paragraph_format.space_after = Pt(2)
            tc_run = tc_p.add_run("\u2714 Accept Terms & Conditions")
            tc_run.bold, tc_run.font.size = True, Pt(11)
            tc_run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)
            tc_desc = doc.add_paragraph(
                "I accept and agree that all information provided in this Job Description is accurate "
                "and complete to the best of my knowledge. I acknowledge that the roles, responsibilities, "
                "qualifications, and company details mentioned in this JD comply with organizational hiring "
                "policies and applicable regulations. I also understand and consent that this Job Description "
                "may be processed, stored, and shared internally for recruitment and evaluation purposes."
            )
            tc_desc.paragraph_format.space_after = Pt(6)
            for run in tc_desc.runs:
                run.font.size = Pt(9)

        if jd_data.get("candidate_name") or jd_data.get("digital_signature_url"):
            doc.add_paragraph()

            sig_head  = doc.add_paragraph()
            sig_head.paragraph_format.space_before = Pt(14)
            sig_head.paragraph_format.space_after  = Pt(12)
            h_run     = sig_head.add_run("Digital Signature Details")
            h_run.bold, h_run.font.size = True, Pt(13)
            h_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

            def _sig_border(cell, color_hex="e2e8f0"):
                tcPr = cell._tc.get_or_add_tcPr()
                bdr  = OxmlElement('w:tcBorders')
                for side in ('top', 'left', 'bottom', 'right'):
                    b = OxmlElement(f'w:{side}')
                    b.set(qn('w:val'),   'single')
                    b.set(qn('w:sz'),    '4')
                    b.set(qn('w:space'), '0')
                    b.set(qn('w:color'), color_hex)
                    bdr.append(b)
                tcPr.append(bdr)

            def _sig_margins(cell, top=200, left=200, bottom=200, right=200):
                tcPr = cell._tc.get_or_get_tcPr()
                mar  = OxmlElement('w:tcMar')
                for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
                    m = OxmlElement(f'w:{side}')
                    m.set(qn('w:w'),    str(val))
                    m.set(qn('w:type'), 'dxa')
                    mar.append(m)
                tcPr.append(mar)

            def _force_script(run):
                rPr = run._element.get_or_add_rPr()
                for existing in rPr.findall(qn('w:rFonts')):
                    rPr.remove(existing)
                rF = OxmlElement('w:rFonts')
                rF.set(qn('w:ascii'), 'Segoe Script')
                rF.set(qn('w:hAnsi'), 'Segoe Script')
                rPr.insert(0, rF)

            sig_tbl = doc.add_table(rows=1, cols=3)
            sig_tbl.autofit = False
            tbl_pr = sig_tbl._tbl.tblPr
            tbl_borders = OxmlElement('w:tblBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'),   'none')
                b.set(qn('w:sz'),    '0')
                b.set(qn('w:color'), 'auto')
                tbl_borders.append(b)
            tbl_pr.append(tbl_borders)
            sig_tbl.columns[0].width = Inches(3.1)
            sig_tbl.columns[1].width = Inches(0.1)
            sig_tbl.columns[2].width = Inches(3.1)

            c_left = sig_tbl.rows[0].cells[0]
            _sig_border(c_left, "e2e8f0")
            set_bg(c_left, "f8fafc")
            _sig_margins(c_left, top=200, left=200, bottom=200, right=200)

            p_badge  = c_left.paragraphs[0]
            p_badge.paragraph_format.space_after = Pt(10)
            badge_run = p_badge.add_run("SIGNEE")
            badge_run.font.size, badge_run.bold = Pt(7.5), True
            badge_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

            inner_tbl = c_left.add_table(rows=1, cols=2)
            inner_tbl.autofit = False
            inner_tbl.columns[0].width = Inches(0.55)
            inner_tbl.columns[1].width = Inches(2.2)

            c_avatar = inner_tbl.rows[0].cells[0]
            set_bg(c_avatar, "eef2ff")
            _sig_margins(c_avatar, top=60, left=0, bottom=0, right=0)
            p_icon = c_avatar.paragraphs[0]
            p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
            icon_run = p_icon.add_run("\U0001F464")
            icon_run.font.size = Pt(20)
            icon_run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)

            c_info = inner_tbl.rows[0].cells[1]
            _sig_margins(c_info, top=30, left=140, bottom=0, right=0)
            p_nm  = c_info.paragraphs[0]
            p_nm.paragraph_format.space_after = Pt(1)
            nm_run = p_nm.add_run(str(jd_data.get("candidate_name") or "Candidate"))
            nm_run.bold, nm_run.font.size = True, Pt(11)
            nm_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

            p_rl   = c_info.add_paragraph()
            rl_run = p_rl.add_run(str(jd_data.get("candidate_role") or "Employee"))
            rl_run.font.size = Pt(8.5)
            rl_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

            c_right = sig_tbl.rows[0].cells[2]
            _sig_border(c_right, "e2e8f0")
            set_bg(c_right, "f8fafc")
            _sig_margins(c_right, top=200, left=180, bottom=200, right=180)

            p_sig = c_right.paragraphs[0]
            p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sig.paragraph_format.space_before = Pt(6)
            p_sig.paragraph_format.space_after  = Pt(8)

            sig_url   = jd_data.get("digital_signature_url")
            sig_name  = str(jd_data.get("candidate_name") or "Candidate")
            sig_added = False
            if sig_url:
                rel_path     = sig_url.lstrip("/")
                abs_sig_path = os.path.join(os.getcwd(), rel_path)
                if os.path.isfile(abs_sig_path):
                    try:
                        p_sig.add_run().add_picture(abs_sig_path, width=Inches(1.8))
                        sig_added = True
                    except Exception:
                        pass
            if not sig_added:
                sig_run = p_sig.add_run(sig_name)
                sig_run.font.size = Pt(22)
                sig_run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
                _force_script(sig_run)
            p_ver = c_right.add_paragraph()
            p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ver_icon = p_ver.add_run("\u2714 ")
            ver_icon.bold, ver_icon.font.size = True, Pt(8)
            ver_icon.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)
            ver_txt = p_ver.add_run("ELECTRONICALLY VERIFIED")
            ver_txt.bold, ver_txt.font.size = True, Pt(7.5)
            ver_txt.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

        buffer_out = io.BytesIO()
        doc.save(buffer_out)
        buffer_out.seek(0)
        return buffer_out.getvalue()

    async def generate_word_stream(self, jd_data: dict, jd_title: str, exclude_terms: bool = False):
        try:
            if not jd_data:
                logger.error("Word export failed: jd_data is None or empty")
                raise ValueError("JD data is empty")
            
            if not jd_title:
                logger.warning("Word export: jd_title is empty, using default")
            
            content = jd_data.get("content", {})
            if not content:
                logger.warning("Word export: content dict is empty for title=%s", jd_title)
            else:
                logger.debug("Word export: content keys = %s", list(content.keys()))
            
            word_bytes = await asyncio.to_thread(self._render_word_bytes, jd_data, jd_title, exclude_terms)
            if not word_bytes:
                logger.error("Word export failed: _render_word_bytes returned empty bytes")
                raise ValueError("Word generation produced no bytes")
            
            logger.debug("Word export: generated %d bytes for title=%s", len(word_bytes), jd_title)
            
            role  = (jd_data.get("job_level") or jd_data.get("seniority") or "role")
            role  = role.replace(" ", "_") if role else "role"
            filename = self._make_attachment_filename(jd_title or "job_description", role, None, "docx")
            logger.debug("Word export: filename=%s", filename)
            
            return StreamingResponse(io.BytesIO(word_bytes),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=\"{filename}\""})
        except Exception as e:
            logger.error(f"Word generation failed: {type(e).__name__}: {str(e)}", exc_info=True)
            raise


pdf_generator = PDFGenerator()
