from __future__ import annotations

import io
import logging
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

logger = logging.getLogger("app.document_ingestion.office")

# Minimum extracted chars relative to file size for legacy .doc heuristic quality check.
_MIN_DOC_CHAR_RATIO = 0.001
_MIN_DOC_CHARS_ABSOLUTE = 20


def extract_office_document(content: bytes, filename: str, fmt: str) -> tuple[str, list[str]]:
    if fmt in {"docx", "word"}:
        return extract_docx(content, filename)
    if fmt == "doc":
        return extract_doc(content, filename)
    if fmt == "rtf":
        return extract_rtf(content, filename)
    raise ValueError(f"No office extractor for format '{fmt}'")


def extract_docx(content: bytes, filename: str) -> tuple[str, list[str]]:
    """Extract text from OOXML (.docx / .word) using python-docx."""
    warnings: list[str] = []
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
    except ImportError as exc:
        raise RuntimeError("python-docx is required: pip install python-docx") from exc

    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        logger.error("Failed to open DOCX %s: %s", filename, exc, exc_info=True)
        raise ValueError(f"Corrupt or invalid DOCX file '{filename}': {exc}") from exc

    parts: list[str] = []
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            text = Paragraph(child, doc).text.strip()
            if text:
                parts.append(text)
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

    if not parts:
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

    text = "\n".join(parts)
    if not text.strip():
        raise ValueError(f"No extractable text found in DOCX: {filename}")

    logger.info("DOCX extraction complete: file=%s chars=%s blocks=%s", filename, len(text), len(parts))
    return text, warnings


def extract_rtf(content: bytes, filename: str) -> tuple[str, list[str]]:
    """Extract text from RTF using striprtf."""
    warnings: list[str] = []
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError as exc:
        raise RuntimeError("striprtf is required: pip install striprtf") from exc

    rtf_text, encoding = _decode_rtf_bytes(content)
    try:
        text = rtf_to_text(rtf_text, encoding=encoding, errors="replace")
    except Exception as exc:
        logger.error("striprtf failed for %s: %s", filename, exc, exc_info=True)
        raise ValueError(f"Failed to parse RTF file '{filename}': {exc}") from exc

    if not text.strip():
        raise ValueError(f"No extractable text found in RTF: {filename}")

    logger.info("RTF extraction complete: file=%s chars=%s", filename, len(text.strip()))
    return text.strip(), warnings


def extract_doc(content: bytes, filename: str) -> tuple[str, list[str]]:
    """
    Extract text from legacy binary .doc (OLE).
    Mislabeled OOXML (.doc extension, PK header) is routed to DOCX extraction.
    """
    warnings: list[str] = []

    if content.startswith(b"PK"):
        warnings.append("File has .doc extension but OOXML (ZIP) content; using DOCX extractor")
        text, docx_w = extract_docx(content, filename)
        warnings.extend(docx_w)
        return text, warnings

    antiword_text = _extract_with_antiword(content, filename)
    if antiword_text:
        warnings.append("Extracted using antiword")
        return antiword_text, warnings

    ole_text = _extract_with_olefile(content, filename)
    if ole_text:
        warnings.append(
            "Extracted using olefile stream heuristic (antiword unavailable). "
            "Verify output quality or re-save as .docx."
        )
        return ole_text, warnings

    fallback = _extract_legacy_binary_strings(content)
    if fallback and len(fallback.strip()) >= _MIN_DOC_CHARS_ABSOLUTE:
        min_expected = max(_MIN_DOC_CHARS_ABSOLUTE, int(len(content) * _MIN_DOC_CHAR_RATIO))
        if len(fallback.strip()) < min_expected:
            raise ValueError(
                f"Legacy .doc extraction for '{filename}' produced insufficient text "
                f"({len(fallback.strip())} chars from {len(content)} bytes). "
                "Re-save as .docx or install antiword for reliable extraction."
            )
        warnings.append(
            "Extracted using binary string fallback (lowest quality). "
            "Re-save as .docx for best results."
        )
        return fallback, warnings

    raise ValueError(
        f"Unable to extract text from legacy .doc file '{filename}'. "
        "Install antiword or re-save the document as .docx."
    )


def _decode_rtf_bytes(content: bytes) -> tuple[str, str]:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return content.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return content.decode("latin-1", errors="replace"), "latin-1"


def _extract_with_antiword(content: bytes, filename: str) -> str | None:
    antiword = shutil.which("antiword")
    if not antiword:
        logger.debug("antiword not found on PATH; skipping for %s", filename)
        return None

    suffix = Path(filename).suffix or ".doc"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        result = subprocess.run(
            [antiword, "-w", "0", tmp_path],
            capture_output=True,
            text=True,
            timeout=30,
            check=False,
        )
        if result.returncode != 0:
            logger.warning(
                "antiword failed for %s: rc=%s stderr=%s",
                filename,
                result.returncode,
                result.stderr.strip(),
            )
            return None
        text = (result.stdout or "").strip()
        return text or None
    except subprocess.TimeoutExpired:
        logger.error("antiword timed out for %s", filename)
        return None
    finally:
        Path(tmp_path).unlink(missing_ok=True)


def _extract_with_olefile(content: bytes, filename: str) -> str | None:
    try:
        import olefile
    except ImportError:
        logger.debug("olefile not installed; skipping OLE extraction for %s", filename)
        return None

    if not olefile.isOleFile(io.BytesIO(content)):
        return None

    try:
        ole = olefile.OleFileIO(io.BytesIO(content))
        parts: list[str] = []
        for stream_name in ("WordDocument", "1Table", "0Table"):
            if ole.exists(stream_name):
                data = ole.openstream(stream_name).read()
                parts.extend(_extract_strings_from_binary(data))
        ole.close()
        text = "\n".join(dict.fromkeys(p for p in parts if p))
        return text.strip() or None
    except Exception as exc:
        logger.warning("olefile extraction failed for %s: %s", filename, exc)
        return None


def _extract_legacy_binary_strings(content: bytes) -> str:
    parts = _extract_strings_from_binary(content)
    return "\n".join(dict.fromkeys(parts))


def _extract_strings_from_binary(data: bytes) -> list[str]:
    parts: list[str] = []
    for match in re.finditer(rb"(?:[\x09\x0a\x0d\x20-\x7e]\x00){4,}", data):
        try:
            decoded = match.group().decode("utf-16-le").strip()
            if len(decoded) >= 4:
                parts.append(decoded)
        except UnicodeDecodeError:
            continue
    for match in re.finditer(rb"[\x09\x0a\x0d\x20-\x7e]{6,}", data):
        decoded = match.group().decode("ascii", errors="ignore").strip()
        if len(decoded) >= 6:
            parts.append(decoded)
    return parts
